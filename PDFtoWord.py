import streamlit as st
import fitz  # PyMuPDF
from io import BytesIO
from docx import Document
from docx.shared import Inches
import tempfile
import os
from PIL import Image

# ✅ Render PDF pages as high-quality images
def render_pdf_pages_as_images(pdf_path, start_page, end_page, zoom=2.0):
    pdf_doc = fitz.open(pdf_path)
    total_pages = len(pdf_doc)
    end_page = min(end_page, total_pages)

    rendered_pages = []
    progress = st.progress(0)

    for i in range(start_page - 1, end_page):
        page = pdf_doc[i]
        mat = fitz.Matrix(zoom, zoom)  # zoom >1 = higher resolution
        pix = page.get_pixmap(matrix=mat)
        img_bytes = pix.tobytes("png")
        rendered_pages.append((i + 1, img_bytes))
        progress.progress((i - (start_page - 1) + 1) / (end_page - start_page + 1))

    return rendered_pages, total_pages

# ✅ Insert page images into Word document safely (no PermissionError)
def generate_word_from_page_images(rendered_pages):
    doc = Document()
    doc.add_heading("PDF as Images (Preserves Layout Exactly)", level=1)

    for page_num, img_data in rendered_pages:
        doc.add_heading(f"Page {page_num}", level=2)

        # Convert raw bytes into a PIL Image
        img = Image.open(BytesIO(img_data))

        # Save a temporary image for python-docx
        tmp_img_path = os.path.join(tempfile.gettempdir(), f"page_{page_num}.png")
        img.save(tmp_img_path, format="PNG")

        # Add the image into the Word document
        width_inches = min(img.width / 300, 6)  # scale width ~6 inches
        doc.add_picture(tmp_img_path, width=Inches(width_inches))

        # Add a page break after each PDF page
        doc.add_page_break()

    # Save final Word document into memory
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ✅ STREAMLIT UI
st.set_page_config(page_title="PDF → Word Exact Layout", layout="wide")
st.title("📄 PDF → Word Converter (Preserve Exact Layout & Images)")

uploaded_file = st.file_uploader("Upload a PDF file", type=["pdf"])

if uploaded_file:
    # Save uploaded PDF to a temporary file
    temp_pdf = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
    temp_pdf.write(uploaded_file.read())
    temp_pdf.close()

    doc = fitz.open(temp_pdf.name)
    total_pages = len(doc)
    st.success(f"✅ PDF has {total_pages} pages")

    # Select page range dynamically
    page_range = st.slider(
        "Select the page range to render",
        min_value=1,
        max_value=total_pages,
        value=(1, min(5, total_pages))  # default: first 5 pages
    )
    start_page, end_page = page_range

    # Select zoom level for better image quality
    zoom_level = st.slider("Zoom Quality (1.0 = normal, 2.0 = HD)", 1.0, 3.0, 2.0)

    if st.button(f"Render Pages {start_page}-{end_page}"):
        st.info(f"Rendering pages {start_page} to {end_page} as high-resolution images...")

        # Render the selected PDF pages into images
        rendered_pages, _ = render_pdf_pages_as_images(
            temp_pdf.name, start_page, end_page, zoom=zoom_level
        )
        st.success(f"✅ Rendered pages {start_page}-{end_page}")

        # Preview first rendered page
        if rendered_pages:
            st.image(rendered_pages[0][1], caption=f"Preview: Page {rendered_pages[0][0]}")

        # Generate Word with exact layout
        file_buffer = generate_word_from_page_images(rendered_pages)
        st.download_button(
            label=f"⬇️ Download Word (Pages {start_page}-{end_page}, exact layout)",
            data=file_buffer,
            file_name=f"pdf_pages_{start_page}_{end_page}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    st.caption("ℹ️ This preserves **exact PDF layout** by rendering each page as an image. "
               "No black images and correct positioning guaranteed.")
